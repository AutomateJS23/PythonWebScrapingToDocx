from bs4 import BeautifulSoup 
from PIL import Image
import requests
import re
import io
from docx import Document
from docx.shared import Inches,Pt
#from docx.enum.section import WD_ORIENT


def headNew(ul): 
    url = requests.get(ul)
    s = BeautifulSoup(url.text, 'lxml')

    headline_main = s.find('div',{'class':'headline_main'})
    a_tag = headline_main.find_all('a')

    contents =[]
    doc = Document()

    for e in a_tag:
        src = e['href']
        prefix_url = requests.get(src)
        prefix_soup = BeautifulSoup(prefix_url.text, 'lxml')
        post_desc = prefix_soup.find('div',{'class':'post_desc'})
        
        detail_text = post_desc.getText()
        text = e.getText()
        text_0 = text.replace("\n","")
        text_1 = text_0.replace("[Clip] ▶▶ ","")
    

        text_sub = detail_text.replace("=","")
        text_sub_0 = text_sub.replace("-","")
        text_n = text_sub_0.replace("\n","")
        text_run = text_n.replace("▶▶","")
    
        contents.append({'SUBJECT': text_1,'DETAIL': text_run, 'LINK': src})

    h = doc.add_heading('Head New',0)
    t = doc.add_table(rows = 1, cols = 3)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = 'SUBJECT'
    t.rows[0].cells[1].text = 'DETAIL'
    t.rows[0].cells[2].text = 'LINK'
    for m in contents:
        fields = t.add_row().cells
        subject_field = fields[0].add_paragraph('').add_run(str(m.get('SUBJECT'))).bold = True
        fields[1].text = str(m.get('DETAIL'))
        fields[2].text = str(m.get('LINK'))
    

    doc.save('Head New.docx')
    
def footballNews(ul):
    url = requests.get(ul)
    s = BeautifulSoup(url.text, 'lxml')

    lastpane = s.find('div',{'class':'lastpanel1'})
    lastNew = lastpane.find_all('div',{'class':'latestnews_tr'})

    contents = []
    doc = Document()
    h = doc.add_heading('Head Line',0)
    t = doc.add_table(rows = 1, cols = 3)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = 'HEADLINE'
    t.rows[0].cells[1].text = 'DRTAIL'
    t.rows[0].cells[2].text = 'image'

    a = 0
    
    for e in lastNew[0:20]:
        a_tag = e.find('a')
        src = a_tag['href']
        
        r = requests.get(src)
        bs_2 = BeautifulSoup(r.text, 'lxml')
        
        tag_title = bs_2.find('div',{'class':'post_head_topic_news'})
        title_text = tag_title.getText()
        
        tag_detail = bs_2.find('div',{'class':'post_desc'})
        tag_detail_2 = tag_detail.find('span', {'style':'font-size: 18px;'})
        
        text = tag_detail_2.getText()
        text_1 = text.replace("\n","")
        
        tag_img = tag_detail.find_all('img')
        a = a + 1 
        for i in tag_img[2:3]:
            name = str(a)
            image = i['src']
            r_img = requests.get(image)
            img_link = Image.open(io.BytesIO(r_img.content))
            img_link.save(name+'.jpg')
            
        contents.append({'Head Line':title_text, 'Deatail':text_1,'image':name+'.jpg'})
        count = len(contents)
        print(f'จำนวนข่าว...{count}')
        
    for m in contents:
        fields = t.add_row().cells
        subject_field = fields[0].add_paragraph('').add_run(str(m.get('Head Line'))).bold = True
        fields[1].text = str(m.get('Deatail'))
        image_field = fields[2].add_paragraph('').add_run().add_picture(str(m.get('image')),width=Inches(2.0))
            
    doc.save('Head Line.docx')
    
def main():
    ul = 'http://www.soccersuck.com/'
    headNew(ul)
    footballNews(ul)



if __name__=='__main__':
    main()
    
