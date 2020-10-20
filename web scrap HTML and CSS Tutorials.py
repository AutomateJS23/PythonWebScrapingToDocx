from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
import requests
import threading



def web_scap(ul,prefix,doc):
    r = requests.get(ul)
    s = BeautifulSoup(r.text,'lxml')
    title = s.find('title')
    cl = s.find_all('li',{'class':'list-group-item'})
    title_text = title.getText()
    h = doc.add_heading(title_text)

    contents = []
    
    for e in cl:
        src = e.find('a')
        link_src = prefix +src['href']
        contents.append({'Subject': src.getText(), 'Link': link_src})
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = 'Subject'
    t.rows[0].cells[1].text = 'Link'

    for m in contents:
        fields = t.add_row().cells
        fields[0].text = str(m.get('Subject'))
        fields[1].text = str(m.get('Link'))


    

def Main():
    doc = Document()
    
    ul = "http://www.java2s.com/Tutorials/HTML_CSS/CSS_Layout_How_to/index.htm"
    prefix = "http://www.java2s.com/Tutorials/HTML_CSS/CSS_Layout_How_to/"
    ul_1 = "http://www.java2s.com/Tutorials/HTML_CSS/Page_Widget_How_to/index.htm"
    prefix_1 = "http://www.java2s.com/Tutorials/HTML_CSS/Page_Widget_How_to/"
    ul_2 = "http://www.java2s.com/Tutorials/HTML_CSS/HTML_Element_Style_How_to/index.htm"
    prefix_2 = "http://www.java2s.com/Tutorials/HTML_CSS/HTML_Element_Style_How_to/"
    ul_3 = "http://www.java2s.com/Tutorials/HTML_CSS/HTML_Form_How_to/index.htm"
    prefix_3 = "http://www.java2s.com/Tutorials/HTML_CSS/HTML_Form_How_to/"
##    web_scap(ul,prefix,doc)
##    web_scap(ul_1,prefix_1,doc)
##    web_scap(ul_2,prefix_2,doc)
##    web_scap(ul_3,prefix_3,doc)

    threads = []

    x = threading.Thread(target=web_scap,args=(ul,prefix,doc))
    x.start()
    threads.append(x)
    x = threading.Thread(target=web_scap,args=(ul_1,prefix_1,doc))
    x.start()
    threads.append(x)
    x = threading.Thread(target=web_scap,args=(ul_2,prefix_2,doc))
    x.start()
    threads.append(x)
    x = threading.Thread(target=web_scap,args=(ul_3,prefix_3,doc))
    x.start()
    threads.append(x)

    for t in threads:
        t.join()
    
    doc.save('Demo_func_1.docx')

    
if __name__=='__main__':
    Main()




