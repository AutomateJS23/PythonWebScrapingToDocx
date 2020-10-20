from bs4 import BeautifulSoup
from PIL import Image
import io
import requests
import re
from docx import Document
#from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt, RGBColor

ul = "https://www.pedigree.com.au/dog-breed-information/dog-breed-gallery"
r = requests.get(ul)
s = BeautifulSoup(r.text,'lxml')

cl = s.find('div',{'id':'dogAZ'})

links = cl.find_all('a')

doc = Document()
h = doc.add_heading('Dog Breeds Gsllery')

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
t.rows[0].cells[0].text = 'Dog name'
t.rows[0].cells[1].text = 'Link'
t.rows[0].cells[2].text = 'image'

contents = []

url = "https://www.pedigree.com.au"
for e in links[:20]:
    img = e.find('img')
    name = img['alt']
    src = url + img['src']
    r = requests.get(src)
    img_link = Image.open(io.BytesIO(r.content))
    img_link.save(e.find('img')['alt']+'.jpg')
    link = url + e['href']
    contents.append({'Dog name': name, 'Link': link, 'image': name+'.jpg'})
    
for m in contents:
    fields = t.add_row().cells
    fields[0].text = str(m.get('Dog name'))
    fields[1].text = str(m.get('Link'))
    image_field = fields[2].add_paragraph('').add_run().add_picture(str(m.get('image')),width=Inches(0.5))
print(contents)


##recodes = [{'Dog name':contents[0][0],'Link':contents[0][1],'image':contents[0][2]},\
##           {'Dog name':contents[1][0],'Link':contents[1][1],'image':contents[1][2]},\
##           {'Dog name':contents[2][0],'Link':contents[2][1],'image':contents[2][2]}]
##    
##for r in recodes:
doc.save('Dog.docx')    
    
    

